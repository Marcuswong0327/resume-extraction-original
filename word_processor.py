import docx
import os
import shutil
import subprocess
import sys
import tempfile
import streamlit as st
from io import BytesIO
from pathlib import Path


class WordProcessor:
    """Handles Word documents (.docx) and legacy Word 97-2003 (.doc) text extraction."""

    def __init__(self):
        pass

    @staticmethod
    def _find_libreoffice_soffice():
        for name in ("soffice", "soffice.exe"):
            path = shutil.which(name)
            if path:
                return path
        if sys.platform == "win32":
            for candidate in (
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ):
                if os.path.isfile(candidate):
                    return candidate
        return None

    def _convert_doc_with_libreoffice(self, doc_bytes: bytes) -> BytesIO | None:
        soffice = self._find_libreoffice_soffice()
        if not soffice:
            return None
        try:
            with tempfile.TemporaryDirectory() as tmp:
                tmp_path = Path(tmp)
                doc_path = tmp_path / "input.doc"
                doc_path.write_bytes(doc_bytes)
                cmd = [
                    soffice,
                    "--headless",
                    "--nologo",
                    "--nofirststartwizard",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(tmp_path),
                    str(doc_path),
                ]
                run_kw: dict = {
                    "check": True,
                    "capture_output": True,
                    "timeout": 120,
                    "text": True,
                }
                if sys.platform == "win32" and hasattr(subprocess, "CREATE_NO_WINDOW"):
                    run_kw["creationflags"] = subprocess.CREATE_NO_WINDOW
                subprocess.run(cmd, **run_kw)
                out_path = tmp_path / "input.docx"
                if out_path.is_file():
                    return BytesIO(out_path.read_bytes())
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError) as e:
            st.warning(f"LibreOffice could not convert .doc: {e}")
        return None

    def _convert_doc_with_win32com(self, doc_bytes: bytes) -> BytesIO | None:
        try:
            import pythoncom
            import win32com.client
        except ImportError:
            return None
        tmpdir = tempfile.mkdtemp()
        try:
            in_path = os.path.abspath(os.path.join(tmpdir, "input.doc"))
            out_path = os.path.abspath(os.path.join(tmpdir, "input.docx"))
            with open(in_path, "wb") as f:
                f.write(doc_bytes)
            pythoncom.CoInitialize()
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                try:
                    word.DisplayAlerts = 0
                except Exception:
                    pass
                doc = word.Documents.Open(
                    in_path,
                    ConfirmConversions=False,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                )
                try:
                    # 16 = wdFormatDocumentDefault (Office Open XML / .docx)
                    doc.SaveAs2(out_path, FileFormat=16)
                finally:
                    doc.Close(SaveChanges=False)
                word.Quit()
            finally:
                pythoncom.CoUninitialize()
            if os.path.isfile(out_path):
                with open(out_path, "rb") as f:
                    return BytesIO(f.read())
        except Exception as e:
            st.warning(f"Microsoft Word could not convert .doc: {e}")
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)
        return None

    def _doc_bytes_to_docx_bytesio(self, doc_bytes: bytes) -> BytesIO | None:
        converted = self._convert_doc_with_libreoffice(doc_bytes)
        if converted is not None:
            return converted
        converted = self._convert_doc_with_win32com(doc_bytes)
        if converted is not None:
            return converted
        st.error(
            "Could not convert .doc to .docx. Install LibreOffice "
            "(https://www.libreoffice.org/) and ensure `soffice` is on PATH, "
            "or install Microsoft Word with pywin32 (`pip install pywin32`), "
            "or re-save the file as .docx."
        )
        return None

    def extract_text_from_docx(self, file_path_or_bytes):
        """
        Extract text from DOCX file.
        Accepts a file path or BytesIO object.
        """
        try:
            doc = docx.Document(file_path_or_bytes)

            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            return "\n".join(text_content)

        except Exception as e:
            st.error(f"Error extracting text from DOCX file: {str(e)}")
            return ""

    def process_word_file(self, uploaded_file):
        """
        Process DOCX or legacy DOC (converted via LibreOffice or Word COM), then extract text.
        """
        try:
            file_extension = uploaded_file.name.lower().split(".")[-1]

            if file_extension == "docx":
                file_content = uploaded_file.read()
                return self.extract_text_from_docx(BytesIO(file_content))

            if file_extension == "doc":
                file_content = uploaded_file.read()
                docx_io = self._doc_bytes_to_docx_bytesio(file_content)
                if docx_io is None:
                    return ""
                docx_io.seek(0)
                return self.extract_text_from_docx(docx_io)

            st.error(f"Unsupported file format: {file_extension}")
            return ""

        except Exception as e:
            st.error(f"Error processing Word file: {str(e)}")
            return ""
